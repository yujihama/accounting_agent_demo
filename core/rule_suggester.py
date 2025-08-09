import json
import tempfile
from typing import Dict, Any, List, Optional
from pathlib import Path
import openai
import logging
from datetime import datetime
import fitz  # PyMuPDF for PDF processing
import docx  # python-docx for Word documents
import pandas as pd
import re
from pydantic import ValidationError
from .models import RuleSuggestionsResponse, RuleSuggestion, EnhancedRuleResponse
import os
from dotenv import load_dotenv

class RuleSuggester:
    """ドキュメントからルール提案を生成するクラス"""
    
    def __init__(self, api_key: str = None):
        # 環境変数をロード
        load_dotenv()
        
        self.api_key = api_key
        self.logger = logging.getLogger(__name__)
        self.provider = os.getenv("OPENAI_PROVIDER", "openai")
        
        # プロバイダーに応じた設定
        if self.provider == "azure":
            self.azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
            self.azure_deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4.1")
            self.azure_api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-15-preview")
    
    def set_api_key(self, api_key: str = None):
        """APIキーを設定"""
        # 引数で渡されたAPIキーか、環境変数から取得
        if api_key:
            self.api_key = api_key
        else:
            if self.provider == "azure":
                self.api_key = os.getenv("AZURE_OPENAI_API_KEY")
            else:
                self.api_key = os.getenv("OPENAI_API_KEY")
        
        if not self.api_key:
            raise ValueError("APIキーが設定されていません。環境変数または引数で指定してください。")
    
    def process_uploaded_document(self, uploaded_file) -> str:
        """アップロードされたドキュメントからテキストを抽出"""
        try:
            file_extension = Path(uploaded_file.name).suffix.lower()
            content = ""
            
            if file_extension == '.pdf':
                content = self._extract_pdf_text(uploaded_file)
            elif file_extension in ['.docx', '.doc']:
                content = self._extract_docx_text(uploaded_file)
            elif file_extension in ['.txt', '.md']:
                content = uploaded_file.read().decode('utf-8')
            elif file_extension in ['.xlsx', '.xls']:
                content = self._extract_excel_text(uploaded_file)
            else:
                raise ValueError(f"サポートされていないファイル形式: {file_extension}")
            
            return content
        except Exception as e:
            self.logger.error(f"ドキュメント処理エラー: {str(e)}")
            raise e
    
    def _extract_pdf_text(self, uploaded_file) -> str:
        """PDFからテキストを抽出"""
        try:
            # BytesIOオブジェクトからPDFを読み込み
            pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            text = ""
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document.load_page(page_num)
                text += page.get_text()
            
            pdf_document.close()
            return text
        except Exception as e:
            raise Exception(f"PDF解析エラー: {str(e)}")
    
    def _extract_docx_text(self, uploaded_file) -> str:
        """Word文書からテキストを抽出"""
        try:
            doc = docx.Document(uploaded_file)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # テーブルのテキストも抽出
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Word文書解析エラー: {str(e)}")
    
    def _extract_excel_text(self, uploaded_file) -> str:
        """Excelファイルからテキストを抽出"""
        try:
            # pandasでExcelファイルを読み込み
            df = pd.read_excel(uploaded_file, sheet_name=None)  # 全シートを読み込み
            text = ""
            
            for sheet_name, sheet_df in df.items():
                text += f"=== {sheet_name} ===\n"
                text += sheet_df.to_string(index=False, na_rep='')
                text += "\n\n"
            
            return text
        except Exception as e:
            raise Exception(f"Excel解析エラー: {str(e)}")
    
    def suggest_rules_from_document(self, document_content: str, existing_rules: Dict[str, Any]) -> List[Dict[str, Any]]:
        """ドキュメント内容から新しいルールを提案"""
        if not self.api_key:
            raise ValueError("OpenAI APIキーが設定されていません")
        
        try:
            # 既存ルールのサマリーを作成
            existing_rules_summary = self._create_existing_rules_summary(existing_rules)
            
            # ルール提案のプロンプトを作成
            prompt = self._create_rule_suggestion_prompt(document_content, existing_rules_summary)
            
            # OpenAI APIを呼び出し (新バージョン対応)
            from openai import OpenAI, AzureOpenAI
            
            # プロバイダーに応じてクライアントを作成
            if self.provider == "azure":
                client = AzureOpenAI(
                    api_key=self.api_key,
                    azure_endpoint=self.azure_endpoint,
                    api_version=self.azure_api_version
                )
                model = self.azure_deployment
            else:
                client = OpenAI(api_key=self.api_key)
                model = "gpt-4.1"
            
            # レスポンススキーマを取得
            response_schema = RuleSuggestionsResponse.model_json_schema()
            
            response = client.chat.completions.create(
                model=model,  # プロバイダーに応じたモデルを使用
                messages=[
                    {
                        "role": "system",
                        "content": f"""あなたは請求書チェックルールの専門家です。
                        提供されたドキュメントを分析し、請求書チェックに有用なルールを提案してください。
                        既存のルールと重複しないよう注意してください。
                        
                        必ず以下のJSONスキーマに厳密に従った形式で回答してください：
                        {json.dumps(response_schema, ensure_ascii=False, indent=2)}"""
                    },
                    {
                        "role": "user", 
                        "content": prompt
                    }
                ],
                temperature=0.3,
                max_tokens=2000,
                response_format={"type": "json_object"}
            )
            
            # レスポンスを解析
            response_text = response.choices[0].message.content.strip()
            
            try:
                # JSONパースとPydantic検証
                json_data = json.loads(response_text)
                validated_response = RuleSuggestionsResponse(**json_data)
                
                # ルールリストを返す
                return [rule.model_dump() for rule in validated_response.suggestions]
                
            except (json.JSONDecodeError, ValidationError) as e:
                self.logger.error(f"ルール提案レスポンス検証エラー: {str(e)}")
                # フォールバックとして従来のパース処理を試みる
                return self._parse_rule_suggestions(response_text)
            
        except Exception as e:
            self.logger.error(f"ルール提案生成エラー: {str(e)}")
            raise e
    
    def _create_existing_rules_summary(self, existing_rules: Dict[str, Any]) -> str:
        """既存ルールのサマリーを作成"""
        if not existing_rules:
            return "既存のルールはありません。"
        
        summary = "既存のルール:\n"
        for rule_id, rule in existing_rules.items():
            summary += f"- {rule['name']} ({rule['category']}): {rule['prompt'][:100]}...\n"
        
        return summary
    
    def _create_rule_suggestion_prompt(self, document_content: str, existing_rules_summary: str) -> str:
        """ルール提案用のプロンプトを作成"""
        return f"""
        以下のドキュメントから、請求書チェックに有用なルールを提案してください。

        【分析対象ドキュメント】
        {document_content[:3000]}  # 長すぎる場合は切り詰め

        【既存のルール】
        {existing_rules_summary}

        【要求事項】
        1. ドキュメントの内容から、請求書チェックに関連する重要な要件を抽出してください
        2. 既存のルールでカバーされていない新しいルールのみを提案してください
        3. 各ルールには適切なカテゴリを設定してください
        4. ルールは具体的で実装可能な内容にしてください

        【出力形式】
        以下のJSON形式で、最大5つのルールを提案してください：

        {{
            "suggested_rules": [
                {{
                    "name": "ルール名",
                    "category": "カテゴリ（日付チェック、金額チェック、承認チェック、書式チェック、その他）",
                    "prompt": "具体的なチェック内容の説明",
                    "confidence": 0.8,
                    "reason": "このルールを提案する理由"
                }}
            ]
        }}
        """
    
    def _parse_rule_suggestions(self, response_text: str) -> List[Dict[str, Any]]:
        """APIレスポンスからルール提案を解析"""
        try:
            # JSONの部分を抽出
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                json_text = json_match.group()
                data = json.loads(json_text)
                return data.get("suggested_rules", [])
            else:
                # JSONが見つからない場合は空のリストを返す
                self.logger.warning("APIレスポンスからJSONを抽出できませんでした")
                return []
        except json.JSONDecodeError as e:
            self.logger.error(f"JSON解析エラー: {str(e)}")
            return []
        except Exception as e:
            self.logger.error(f"ルール提案解析エラー: {str(e)}")
            return []
    
    def validate_suggested_rule(self, rule_data: Dict[str, Any]) -> List[str]:
        """提案されたルールの妥当性を検証"""
        errors = []
        
        required_fields = ["name", "category", "prompt"]
        for field in required_fields:
            if not rule_data.get(field):
                errors.append(f"必須フィールド '{field}' が不足しています")
        
        if rule_data.get("name") and len(rule_data["name"]) < 3:
            errors.append("ルール名は3文字以上である必要があります")
        
        if rule_data.get("prompt") and len(rule_data["prompt"]) < 10:
            errors.append("プロンプトは10文字以上である必要があります")
        
        valid_categories = ["日付チェック", "金額チェック", "承認チェック", "書式チェック", "その他"]
        if rule_data.get("category") and rule_data["category"] not in valid_categories:
            errors.append(f"無効なカテゴリです。有効なカテゴリ: {', '.join(valid_categories)}")
        
        return errors
    
    def enhance_rule_suggestion(self, base_rule: Dict[str, Any], document_content: str) -> Dict[str, Any]:
        """ドキュメント内容を基に、提案されたルールをさらに詳細化"""
        if not self.api_key:
            return base_rule
        
        try:
            prompt = f"""
            以下のルール提案をドキュメント内容に基づいてより詳細で具体的にしてください：

            【基本ルール】
            名前: {base_rule['name']}
            カテゴリ: {base_rule['category']}
            内容: {base_rule['prompt']}

            【参考ドキュメント】
            {document_content[:2000]}

            【要求】
            ルールの内容をより具体的で実装可能な形に改善してください。
            チェック項目を明確にし、判定基準を具体的に記述してください。

            以下のJSON形式で出力してください：
            {{
                "enhanced_prompt": "改善されたルール内容"
            }}
            """
            
            from openai import OpenAI, AzureOpenAI
            
            # プロバイダーに応じてクライアントを作成
            if self.provider == "azure":
                client = AzureOpenAI(
                    api_key=self.api_key,
                    azure_endpoint=self.azure_endpoint,
                    api_version=self.azure_api_version
                )
                model = self.azure_deployment
            else:
                client = OpenAI(api_key=self.api_key)
                model = "gpt-4.1"
            
            # レスポンススキーマを取得
            enhanced_schema = EnhancedRuleResponse.model_json_schema()
            
            response = client.chat.completions.create(
                model=model,  # プロバイダーに応じたモデルを使用
                messages=[
                    {
                        "role": "system",
                        "content": f"""あなたは請求書チェックルールの改善専門家です。
                        
                        必ず以下のJSONスキーマに従った形式で回答してください：
                        {json.dumps(enhanced_schema, ensure_ascii=False, indent=2)}"""
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.2,
                max_tokens=1000,
                response_format={"type": "json_object"}
            )
            
            response_text = response.choices[0].message.content.strip()
            
            try:
                # JSONパースとPydantic検証
                json_data = json.loads(response_text)
                validated_response = EnhancedRuleResponse(**json_data)
                
                enhanced_rule = base_rule.copy()
                enhanced_rule['prompt'] = validated_response.enhanced_prompt
                
                # 改善点があれば追加
                if validated_response.improvements:
                    enhanced_rule['improvements'] = validated_response.improvements
                
                return enhanced_rule
                
            except (json.JSONDecodeError, ValidationError) as e:
                self.logger.error(f"ルール改善レスポンス検証エラー: {str(e)}")
                # フォールバック：従来のパース処理
                json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                if json_match:
                    try:
                        data = json.loads(json_match.group())
                        enhanced_rule = base_rule.copy()
                        enhanced_rule['prompt'] = data.get('enhanced_prompt', base_rule['prompt'])
                        return enhanced_rule
                    except:
                        pass
            
        except Exception as e:
            self.logger.error(f"ルール改善エラー: {str(e)}")
        
        return base_rule