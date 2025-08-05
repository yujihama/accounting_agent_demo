import io
import tempfile
from pathlib import Path
from typing import Dict, List, Any, Optional
import pandas as pd
from datetime import datetime

# PDFファイル処理用
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Wordファイル処理用
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class FileProcessor:
    """ファイル処理機能を提供するクラス"""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_excel,
            'application/vnd.ms-excel': self._process_excel,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/msword': self._process_doc,
        }
    
    def process_files(self, uploaded_files) -> Dict[str, Dict[str, Any]]:
        """
        アップロードされたファイルを処理
        
        Args:
            uploaded_files: Streamlitのアップロードファイルオブジェクトのリスト
            
        Returns:
            ファイル名をキーとした処理結果の辞書
        """
        results = {}
        
        for file in uploaded_files:
            try:
                result = self._process_single_file(file)
                results[file.name] = result
            except Exception as e:
                results[file.name] = {
                    "file_name": file.name,
                    "error": f"ファイル処理エラー: {str(e)}",
                    "metadata": {
                        "file_type": file.type,
                        "file_size": file.size,
                        "processed_at": datetime.now().isoformat()
                    }
                }
        
        return results
    
    def _process_single_file(self, file) -> Dict[str, Any]:
        """単一ファイルを処理"""
        file_type = file.type
        file_name = file.name
        
        # ファイルの内容を読み取り
        file_content = file.read()
        file.seek(0)  # ファイルポインタをリセット
        
        # ファイルタイプに応じて処理
        if file_type in self.supported_types:
            processor = self.supported_types[file_type]
            content = processor(file_content, file_name)
        else:
            # サポートされていないファイルタイプの場合、拡張子で判定
            extension = Path(file_name).suffix.lower()
            if extension == '.pdf':
                content = self._process_pdf(file_content, file_name)
            elif extension in ['.xlsx', '.xls']:
                content = self._process_excel_by_extension(file_content, file_name, extension)
            elif extension == '.docx':
                content = self._process_docx(file_content, file_name)
            else:
                content = f"サポートされていないファイル形式: {file_type} ({extension})"
        
        return {
            "file_name": file_name,
            "content": content,
            "metadata": {
                "file_type": file_type,
                "file_size": file.size,
                "processed_at": datetime.now().isoformat(),
                "extension": Path(file_name).suffix.lower()
            }
        }
    
    def _process_pdf(self, file_content: bytes, file_name: str) -> str:
        """PDFファイルを処理"""
        if not PDF_AVAILABLE:
            return "PDFファイルの処理にはPyPDF2が必要です"
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"--- ページ {page_num + 1} ---")
                    text_content.append(page_text)
            
            if not text_content:
                return "PDFからテキストを抽出できませんでした"
            
            return "\n".join(text_content)
            
        except Exception as e:
            return f"PDF処理エラー: {str(e)}"
    
    def _process_excel(self, file_content: bytes, file_name: str) -> str:
        """Excelファイルを処理"""
        try:
            excel_file = io.BytesIO(file_content)
            
            # すべてのシートを読み込み
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            
            text_content = []
            for sheet_name, df in excel_data.items():
                text_content.append(f"--- シート: {sheet_name} ---")
                
                # データフレームを文字列に変換
                df_string = df.to_string(index=False, na_rep='')
                text_content.append(df_string)
                text_content.append("")
            
            return "\n".join(text_content)
            
        except Exception as e:
            return f"Excel処理エラー: {str(e)}"
    
    def _process_excel_by_extension(self, file_content: bytes, file_name: str, extension: str) -> str:
        """拡張子に基づいてExcelファイルを処理"""
        try:
            excel_file = io.BytesIO(file_content)
            
            if extension == '.xlsx':
                engine = 'openpyxl'
            else:  # .xls
                engine = 'xlrd'
            
            excel_data = pd.read_excel(excel_file, sheet_name=None, engine=engine)
            
            text_content = []
            for sheet_name, df in excel_data.items():
                text_content.append(f"--- シート: {sheet_name} ---")
                df_string = df.to_string(index=False, na_rep='')
                text_content.append(df_string)
                text_content.append("")
            
            return "\n".join(text_content)
            
        except Exception as e:
            return f"Excel処理エラー: {str(e)}"
    
    def _process_docx(self, file_content: bytes, file_name: str) -> str:
        """Word文書を処理"""
        if not DOCX_AVAILABLE:
            return "Word文書の処理にはpython-docxが必要です"
        
        try:
            doc_file = io.BytesIO(file_content)
            document = Document(doc_file)
            
            text_content = []
            
            # 段落を処理
            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # テーブルを処理
            for table in document.tables:
                text_content.append("\n--- テーブル ---")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    text_content.append(" | ".join(row_text))
            
            if not text_content:
                return "Word文書からテキストを抽出できませんでした"
            
            return "\n".join(text_content)
            
        except Exception as e:
            return f"Word文書処理エラー: {str(e)}"
    
    def _process_doc(self, file_content: bytes, file_name: str) -> str:
        """古いWord文書(.doc)を処理"""
        # .docファイルの処理は複雑なため、基本的なメッセージを返す
        return f"旧形式のWord文書 (.doc) はサポートされていません。{file_name} を .docx 形式で保存し直してください。"
    
    def validate_file(self, file) -> tuple[bool, Optional[str]]:
        """
        ファイルの妥当性を検証
        
        Returns:
            (is_valid, error_message)
        """
        # ファイルサイズチェック (50MB以下)
        max_size = 50 * 1024 * 1024  # 50MB
        if file.size > max_size:
            return False, f"ファイルサイズが大きすぎます（最大50MB）: {file.size / 1024 / 1024:.1f}MB"
        
        # ファイル名チェック
        if not file.name:
            return False, "ファイル名が無効です"
        
        # 拡張子チェック
        allowed_extensions = ['.pdf', '.xlsx', '.xls', '.docx', '.doc']
        file_extension = Path(file.name).suffix.lower()
        
        if file_extension not in allowed_extensions:
            return False, f"サポートされていないファイル形式です: {file_extension}"
        
        return True, None
    
    def get_file_info(self, file) -> Dict[str, Any]:
        """ファイル情報を取得"""
        return {
            "name": file.name,
            "size": file.size,
            "type": file.type,
            "extension": Path(file.name).suffix.lower()
        }
    
    def extract_metadata(self, file_content: bytes, file_type: str) -> Dict[str, Any]:
        """ファイルからメタデータを抽出"""
        metadata = {
            "extraction_time": datetime.now().isoformat(),
            "file_type": file_type
        }
        
        try:
            if file_type == 'application/pdf' and PDF_AVAILABLE:
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                metadata.update({
                    "page_count": len(pdf_reader.pages),
                    "pdf_info": pdf_reader.metadata if hasattr(pdf_reader, 'metadata') else {}
                })
            
            elif 'excel' in file_type or 'sheet' in file_type:
                excel_file = io.BytesIO(file_content)
                excel_data = pd.read_excel(excel_file, sheet_name=None)
                metadata.update({
                    "sheet_count": len(excel_data),
                    "sheet_names": list(excel_data.keys())
                })
            
            elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' and DOCX_AVAILABLE:
                doc_file = io.BytesIO(file_content)
                document = Document(doc_file)
                metadata.update({
                    "paragraph_count": len(document.paragraphs),
                    "table_count": len(document.tables)
                })
                
        except Exception as e:
            metadata["metadata_error"] = str(e)
        
        return metadata